
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS_DIR = "e:/coastal/docs"
OUTPUT_DOCX = "e:/coastal/Coastal_Banking_Documentation_Full.docx"

FILES = [
    "index.md",
    "01-architecture.md",
    "02-data-models.md",
    "03-logic-flows.md",
    "04-ui-ux.md",
    "05-security.md",
    "06-infrastructure.md",
    "07-api-reference.md"
]

def add_markdown_content(doc, content):
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        # Code Blocks (Mermaid/Code) - simplified
        elif line.startswith('```'):
            continue # Skip fence lines
        else:
            p = doc.add_paragraph(line)

def create_docx():
    doc = Document()

    # Title Page
    title = doc.add_heading('Coastal Banking System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    for filename in FILES:
        filepath = os.path.join(DOCS_DIR, filename)
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as infile:
                content = infile.read()
                # Skip the YAML header if present
                doc.add_page_break()
                add_markdown_content(doc, content)

    doc.save(OUTPUT_DOCX)
    print(f"Successfully created: {OUTPUT_DOCX}")

if __name__ == "__main__":
    create_docx()
